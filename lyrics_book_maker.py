import os
import re
import glob
import sys
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# 全局字体设置
LXGW_FONT_NAME = "霞鹜新晰黑"
LXGW_FONT_PATH = "fonts/LXGWNeoXiHei.ttf"

def set_double_columns(section):
    """设置双栏布局和栏间距"""
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '8')  # 缩小栏间距
    cols.set(qn('w:equalWidth'), '1')  # 等宽栏
    sectPr.append(cols)

def detect_lrc_type(file_path):
    """检测LRC文件类型（中文、日文或英文）"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 检查是否包含日文字符（平假名和片假名）
        if re.search(r'[\u3040-\u309F\u30A0-\u30FF]', content):
            return 'japanese'
        # 检查是否包含中文字符
        elif re.search(r'[\u4e00-\u9fff]', content):
            return 'chinese'
        # 检查是否包含英文字符
        elif re.search(r'[a-zA-Z]', content):
            return 'english'
        return None
    except Exception as e:
        print(f"检测文件类型时出错: {e}")
        return None

def set_font(run, font_name=LXGW_FONT_NAME):
    """设置字体为霞鹜新晰黑"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

def parse_japanese_lrc(file_path):
    """解析日文LRC文件，返回结构化歌词数据"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        lyrics = []
        current_group = []
        
        for line in lines:
            line = line.strip()
            # 跳过元数据行
            if not line or (line.startswith('[') and ('作词' in line or '作曲' in line or '编曲' in line or 'by:' in line)):
                continue
            
            # 处理歌词行
            if re.match(r'^\[\d+:\d+\.\d+\]', line):
                text = re.sub(r'^\[\d+:\d+\.\d+\]\s*', '', line)
                if text:
                    current_group.append(text)
                    
                    # 每三行组成一组 (日/中/罗马音)
                    if len(current_group) == 3:
                        # 确保顺序正确: 日语、中文、罗马音
                        jp_line = current_group[0]
                        cn_line = current_group[1]
                        roma_line = current_group[2]
                        
                        # 如果顺序不对，尝试自动调整
                        if not re.search(r'[\u3040-\u309F\u30A0-\u30FF]', jp_line):
                            # 如果第一行不是日文，尝试重新排序
                            for i, text in enumerate(current_group):
                                if re.search(r'[\u3040-\u309F\u30A0-\u30FF]', text):
                                    jp_line = text
                                elif re.search(r'[\u4e00-\u9fff]', text):
                                    cn_line = text
                                else:
                                    roma_line = text
                        
                        lyrics.append({
                            'japanese': jp_line,
                            'chinese': cn_line,
                            'romaji': roma_line
                        })
                        current_group = []
        
        # 处理最后一组不完整的歌词（如果有）
        if current_group:
            jp_line = current_group[0] if current_group else ""
            cn_line = current_group[1] if len(current_group) > 1 else ""
            roma_line = current_group[2] if len(current_group) > 2 else ""
            lyrics.append({
                'japanese': jp_line,
                'chinese': cn_line,
                'romaji': roma_line
            })
        
        return lyrics
    except Exception as e:
        print(f"解析日文歌词文件时出错: {e}")
        return []

def parse_chinese_lrc(file_path):
    """解析中文LRC文件，返回结构化歌词数据"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        lyrics = []
        current_lyric = ""
        current_pinyin = ""
        
        for line in lines:
            line = line.strip()
            
            # 跳过元数据行
            if not line or (line.startswith('[') and ('作词' in line or '作曲' in line or '编曲' in line or '制作人' in line or 'by:' in line)):
                continue
            
            # 处理歌词行
            if re.match(r'^\[\d+:\d+\.\d+\]', line):
                # 移除时间标签
                text = re.sub(r'^\[\d+:\d+\.\d+\]\s*', '', line)
                
                if text:
                    # 检测是否为拼音行 (只包含字母、数字和空格)
                    if re.match(r'^[a-zA-Z0-9\s\']+$', text):
                        current_pinyin = text
                    else:
                        # 如果是中文歌词行
                        if current_lyric:
                            # 保存前一组歌词
                            lyrics.append({
                                'lyric': current_lyric,
                                'pinyin': current_pinyin
                            })
                            current_pinyin = ""
                        
                        current_lyric = text
        
        # 添加最后一组歌词
        if current_lyric:
            lyrics.append({
                'lyric': current_lyric,
                'pinyin': current_pinyin
            })
        
        return lyrics
    except Exception as e:
        print(f"解析中文歌词文件时出错: {e}")
        return []

def parse_english_lrc(file_path):
    """解析英文LRC文件，返回结构化歌词数据"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        lyrics = []
        current_english = ""
        current_chinese = ""
        
        for line in lines:
            line = line.strip()
            
            # 跳过元数据行
            if not line or (line.startswith('[') and ('作词' in line or '作曲' in line or '编曲' in line or '制作人' in line or 'by:' in line)):
                continue
            
            # 处理歌词行
            if re.match(r'^\[\d+:\d+\.\d+\]', line):
                # 移除时间标签
                text = re.sub(r'^\[\d+:\d+\.\d+\]\s*', '', line)
                
                if text:
                    # 检测是否为中文行 (包含中文字符)
                    if re.search(r'[\u4e00-\u9fff]', text):
                        current_chinese = text
                        # 保存当前组
                        if current_english or current_chinese:
                            lyrics.append({
                                'english': current_english,
                                'chinese': current_chinese
                            })
                            current_english = ""
                            current_chinese = ""
                    else:
                        # 如果是英文歌词行
                        current_english = text
        
        # 添加最后一组歌词
        if current_english or current_chinese:
            lyrics.append({
                'english': current_english,
                'chinese': current_chinese
            })
        
        return lyrics
    except Exception as e:
        print(f"解析英文歌词文件时出错: {e}")
        return []

def create_page_number_field():
    """创建页码字段"""
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    
    return [fld_char, instr_text, fld_char2]

def create_lyrics_book():
    """创建智能歌词本主函数"""
    try:
        print("="*50)
        print("多语言歌词本生成器 - 支持中日英三种语言")
        print("="*50)
        print("正在检测歌词文件...")
        
        # 创建必要目录
        os.makedirs('lrc_files', exist_ok=True)
        os.makedirs('output', exist_ok=True)
        
        doc = Document()
        
        # ===== 页面设置 =====
        section = doc.sections[0]
        section.page_width = Mm(148)  # A5宽度
        section.page_height = Mm(210)  # A5高度
        
        # 缩小页边距以节省空间
        section.left_margin = Mm(8)
        section.right_margin = Mm(8)
        section.top_margin = Mm(6)
        section.bottom_margin = Mm(6)
        
        # ===== 设置全局字体 =====
        # 设置正文样式
        style = doc.styles['Normal']
        font = style.font
        font.name = LXGW_FONT_NAME
        font._element.rPr.rFonts.set(qn('w:eastAsia'), LXGW_FONT_NAME)
        
        # 设置标题样式
        for level in range(1, 6):
            heading_style = doc.styles[f'Heading {level}']
            heading_font = heading_style.font
            heading_font.name = LXGW_FONT_NAME
            heading_font._element.rPr.rFonts.set(qn('w:eastAsia'), LXGW_FONT_NAME)
        
        # ===== 创建目录 =====
        title = doc.add_heading('多语言歌词本目录', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.size = Pt(10)  # 缩小标题字体
        set_font(title.runs[0])
        
        # 设置目录部分为双栏
        set_double_columns(doc.sections[0])
        
        song_data = []
        lang_count = {'japanese': 0, 'chinese': 0, 'english': 0}
        page_numbers = {}  # 存储歌曲标题对应的页码
        
        # 解析所有歌曲并检测语言类型
        lrc_files = glob.glob('lrc_files/*.lrc')
        
        if not lrc_files:
            print("\n警告: 未找到任何.lrc歌词文件！")
            print("请将歌词文件放入 'lrc_files' 文件夹中")
            return
        
        print(f"找到 {len(lrc_files)} 个歌词文件")
        
        # 添加目录内容
        for idx, file_path in enumerate(lrc_files):
            filename = os.path.basename(file_path)
            song_title = os.path.splitext(filename)[0]
            
            # 检测歌词类型
            lrc_type = detect_lrc_type(file_path)
            if not lrc_type:
                print(f"警告: 无法确定 '{song_title}' 的语言类型，跳过处理")
                continue
                
            # 根据类型解析歌词
            if lrc_type == 'japanese':
                lyrics = parse_japanese_lrc(file_path)
                lang_count['japanese'] += 1
            elif lrc_type == 'chinese':
                lyrics = parse_chinese_lrc(file_path)
                lang_count['chinese'] += 1
            else:  # 英文
                lyrics = parse_english_lrc(file_path)
                lang_count['english'] += 1
            
            song_data.append({
                'title': song_title,
                'file_path': file_path,
                'lyrics': lyrics,
                'type': lrc_type
            })
            
            print(f"处理: {song_title} ({lrc_type}) - {len(lyrics)}行歌词")
            
            # 添加目录项（带页码占位符）
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)  # 减小行距
            p.paragraph_format.space_before = Pt(1)
            
            # 添加歌曲标题
            run_title = p.add_run(f"{idx+1}. {song_title}")
            run_title.font.size = Pt(8)  # 缩小目录字体
            set_font(run_title)
            
            # 添加制表符
            p.add_run("\t")
            
            # 添加页码占位符（实际页码将在生成内容后更新）
            run_page = p.add_run(" ")
            run_page.font.size = Pt(8)
            set_font(run_page)
            # 存储对页码元素的引用，稍后更新
            page_numbers[song_title] = run_page
        
        if not song_data:
            print("错误: 没有找到可处理的歌词文件！")
            return
        
        doc.add_page_break()
        
        # ===== 添加歌词内容 =====
        print("\n正在生成歌词本内容...")
        current_page = 1  # 从第一页开始计数
        
        for song_idx, song in enumerate(song_data):
            if not song['lyrics']:
                continue
            
            # 添加新节（每首歌独立一节）
            if song_idx > 0:
                doc.add_section(WD_SECTION.NEW_PAGE)
                current_page += 1  # 每首歌从新页开始
            
            current_section = doc.sections[-1]
            set_double_columns(current_section)  # 设置双栏
            
            # 记录歌曲起始页码
            song_start_page = current_page
            
            # 歌曲标题（跨栏显示）
            title_para = doc.add_heading(song['title'], level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.style.font.size = Pt(9)  # 缩小标题字体
            set_font(title_para.runs[0])
            
            # 更新目录中的页码
            if song['title'] in page_numbers:
                page_run = page_numbers[song['title']]
                page_run.text = f"第 {song_start_page} 页"
            
            # 根据类型添加歌词内容
            if song['type'] == 'japanese':
                # 日文歌曲处理
                for line in song['lyrics']:
                    # 日语+罗马音行（同一行）
                    p_jp = doc.add_paragraph()
                    p_jp.paragraph_format.space_after = Pt(0)  # 减小段落间距
                    run_jp = p_jp.add_run(line['japanese'])
                    run_jp.font.size = Pt(8)  # 缩小日文字体
                    set_font(run_jp)
                    
                    # 添加罗马音（灰色小字）
                    if line['romaji']:  # 确保有罗马音内容
                        run_roma = p_jp.add_run("  " + line['romaji'])
                        run_roma.font.size = Pt(6.5)  # 缩小罗马音字体
                        run_roma.font.color.rgb = RGBColor(100, 100, 100)
                        set_font(run_roma)
                    
                    # 中文翻译行（使用中文字体）
                    if line['chinese']:  # 确保有中文内容
                        p_cn = doc.add_paragraph()
                        p_cn.paragraph_format.space_after = Pt(3)  # 减小段落间距
                        run_cn = p_cn.add_run(line['chinese'])
                        run_cn.font.size = Pt(7.5)  # 缩小中文字体
                        run_cn.font.color.rgb = RGBColor(50, 50, 150)
                        set_font(run_cn)
            
            elif song['type'] == 'chinese':
                # 中文歌曲处理
                for line in song['lyrics']:
                    # 中文歌词行
                    p_lyric = doc.add_paragraph()
                    p_lyric.paragraph_format.space_after = Pt(0)  # 减小段落间距
                    run_lyric = p_lyric.add_run(line['lyric'])
                    run_lyric.font.size = Pt(8)  # 缩小歌词字体
                    set_font(run_lyric)
                    
                    # 添加拼音行（如果有）
                    if line['pinyin']:
                        p_pinyin = doc.add_paragraph()
                        p_pinyin.paragraph_format.space_after = Pt(3)  # 减小段落间距
                        run_pinyin = p_pinyin.add_run(line['pinyin'])
                        run_pinyin.font.size = Pt(7)  # 缩小拼音字体
                        run_pinyin.font.color.rgb = RGBColor(100, 100, 100)
                        set_font(run_pinyin)
            
            else:  # 英文歌曲处理
                for line in song['lyrics']:
                    # 英文歌词行
                    p_en = doc.add_paragraph()
                    p_en.paragraph_format.space_after = Pt(0)  # 减小段落间距
                    run_en = p_en.add_run(line['english'])
                    run_en.font.size = Pt(8)  # 缩小英文字体
                    set_font(run_en)
                    
                    # 中文翻译行（如果有）
                    if line['chinese']:
                        p_cn = doc.add_paragraph()
                        p_cn.paragraph_format.space_after = Pt(3)  # 减小段落间距
                        run_cn = p_cn.add_run(line['chinese'])
                        run_cn.font.size = Pt(7.5)  # 缩小中文字体
                        run_cn.font.color.rgb = RGBColor(50, 50, 150)
                        set_font(run_cn)
            
            # 歌曲间分隔（不使用分页符，改为紧凑布局）
            doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
            
            # 更新当前页码（每首歌后增加一页）
            current_page += 1
        
        # 保存文档
        output_path = 'output/多语言歌词本.docx'
        doc.save(output_path)
        
        print("\n" + "="*50)
        print(f"成功生成歌词本: {output_path}")
        print(f"包含歌曲: 日文 {lang_count['japanese']} 首, 中文 {lang_count['chinese']} 首, 英文 {lang_count['english']} 首")
        print("="*50)
        print("\n打印提示：")
        print("1. 纸张大小: A5")
        print("2. 每张纸打印页数: 2 (双面打印时更省纸)")
        print("3. 缩放比例: 100%")
        print("\n操作完成！")
    
    except Exception as e:
        print(f"生成歌词本时发生错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    create_lyrics_book()
    
    # 添加等待用户输入，防止窗口关闭
    if getattr(sys, 'frozen', False):  # 打包后运行
        print("\n按Enter键退出...")
        input()